import os
import io
import logging
from typing import Optional, Dict, Any
from PIL import Image
import pytesseract
import PyPDF2
from pathlib import Path
import docx
import time

logger = logging.getLogger(__name__)

# Configure Tesseract path for Windows
if os.name == 'nt':  # Windows
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class OCRProcessor:
    """Process various file types with error handling and CPU optimization"""
    
    SUPPORTED_IMAGE_FORMATS = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
    SUPPORTED_DOC_FORMATS = ['.pdf', '.docx', '.txt', '.xlsx']
    
    # CPU Optimization settings
    MAX_IMAGE_SIZE = (1500, 1500)  # Resize large images for faster OCR
    OCR_TIMEOUT = 30  # seconds
    MIN_CONFIDENCE_THRESHOLD = 30  # Tesseract confidence threshold
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def process_file(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process file with comprehensive error handling"""
        result = {
            'success': False,
            'text': '',
            'error': None,
            'processing_time': 0,
            'fallback_used': False
        }
        
        start_time = time.time()
        file_ext = Path(filename).suffix.lower()
        
        try:
            # Validate file type
            if file_ext not in self.SUPPORTED_IMAGE_FORMATS + self.SUPPORTED_DOC_FORMATS:
                result['error'] = f"Unsupported file format: {file_ext}"
                return result
            
            # Process based on file type
            if file_ext in self.SUPPORTED_IMAGE_FORMATS:
                text = self._extract_text_from_image_optimized(file_bytes)
            elif file_ext == '.pdf':
                text = self._extract_text_from_pdf_optimized(file_bytes)
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_bytes)
            elif file_ext == '.txt':
                text = file_bytes.decode('utf-8', errors='ignore')
            else:
                result['error'] = f"Handler not implemented for {file_ext}"
                return result
            
            # Validate extracted text
            if not text or len(text.strip()) < 5:
                result['error'] = "Could not extract meaningful text from the file"
                result['fallback_used'] = True
                result['text'] = "[Document uploaded but text extraction failed]"
            else:
                result['success'] = True
                result['text'] = text.strip()
                
        except Exception as e:
            self.logger.error(f"OCR processing error for {filename}: {e}")
            result['error'] = str(e)
            result['fallback_used'] = True
            result['text'] = "[Document processing error - using question only]"
            
        result['processing_time'] = time.time() - start_time
        return result
    
    def _extract_text_from_image_optimized(self, image_bytes: bytes) -> str:
        """Optimized OCR for CPU processing"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            
            # Convert to RGB
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # CPU Optimization: Resize large images
            if image.size[0] > self.MAX_IMAGE_SIZE[0] or image.size[1] > self.MAX_IMAGE_SIZE[1]:
                image.thumbnail(self.MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)
                self.logger.info(f"Resized image for faster OCR: {image.size}")
            
            # CPU Optimization: Use faster OCR settings
            custom_config = r'--oem 3 --psm 3 -c tessedit_do_invert=0'
            
            # Extract text with timeout protection
            try:
                text = pytesseract.image_to_string(
                    image,
                    lang='eng',  # Use only English for speed
                    config=custom_config,
                    timeout=self.OCR_TIMEOUT
                )
            except RuntimeError as timeout_error:
                self.logger.error(f"OCR timeout: {timeout_error}")
                return "[OCR timeout - text extraction took too long]"
                
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Image OCR error: {e}")
            raise
    
    def _extract_text_from_pdf_optimized(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF with fallback to OCR only if needed"""
        text_parts = []
        
        try:
            # First try direct text extraction (much faster than OCR)
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            
            for page_num, page in enumerate(pdf_reader.pages[:10]):  # Limit to first 10 pages
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            
            # If we got text, return it (no OCR needed)
            if text_parts and any(t.strip() for t in text_parts):
                return "\n\n".join(text_parts)
            
            # Only use OCR if no text was extracted (scanned PDF)
            self.logger.info("No text in PDF, attempting OCR (this may be slow)...")
            
            # For CPU: Only OCR first 2 pages to save time
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(pdf_bytes, last_page=2, dpi=150)  # Lower DPI for speed
                
                for i, image in enumerate(images):
                    # Resize for faster OCR
                    if image.size[0] > self.MAX_IMAGE_SIZE[0]:
                        image.thumbnail(self.MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)
                    
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    if page_text.strip():
                        text_parts.append(f"Page {i+1}:\n{page_text}")
                        
            except ImportError:
                return "[PDF is scanned - OCR required but pdf2image not installed]"
                
        except Exception as e:
            self.logger.error(f"PDF processing error: {e}")
            raise
            
        return "\n\n".join(text_parts) if text_parts else ""
    
    def _extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs[:100]:  # Limit paragraphs for speed
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract first few tables
            for table in doc.tables[:5]:  # Limit tables
                for row in table.rows[:20]:  # Limit rows
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                        
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"DOCX processing error: {e}")
            raise